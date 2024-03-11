"""
Функции для генерации выходного файла с оформленным списком использованных источников.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # pylint: disable=E0611
from docx.shared import Pt


class Renderer:
    """
    Создание выходного файла – Word.
    """

    def __init__(self, rows: tuple[str, ...]):
        self.rows = rows

    def render(self, path: Path | str) -> None:
        """
        Метод генерации Word-файла со списком использованных источников.

        :param Path | str path: Путь для сохранения выходного файла.
        """

        document = Document()

        # стилизация заголовка
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        runner = paragraph.add_run("Список использованной литературы")
        runner.bold = True

        # стилизация текста
        style_normal = document.styles["Normal"]
        style_normal.font.name = "Times New Roman"
        style_normal.font.size = Pt(12)
        style_normal.paragraph_format.line_spacing = 1.5
        style_normal.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        for row in self.rows:
            # добавление источника
            document.add_paragraph(row, style="List Number")

        # сохранение файла Word
        document.save(path)
